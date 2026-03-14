import os
from dotenv import load_dotenv
import subprocess
import asyncio
from datetime import datetime
from contextlib import asynccontextmanager
from typing import Optional
from agent import ChatAgent
from fastapi import FastAPI, Request, Header, HTTPException, Body
from fastapi.concurrency import run_in_threadpool
from httpx import AsyncClient
from pydantic import BaseModel
from docx import Document

# Load environment variables from .env file
load_dotenv()

# Initialize the agent globally (outside the function) so it persists
agent = ChatAgent()


# Create FastAPI app with lifespan
app = FastAPI()

BOT_TOKEN = os.getenv("BOT_TOKEN")
SECRET_TOKEN = os.getenv("SECRET_TOKEN")
MY_CHAT_ID = os.getenv("MY_CHAT_ID")
TELEGRAM_API_URL = f"https://api.telegram.org/bot{BOT_TOKEN}"
KNOWLEDGE_BASE_PATH = "knowledge_base.docx"


# Pydantic model for feed request body
class FeedRequest(BaseModel):
    text: str


@app.get("/health")
async def health():
    """Health check endpoint for UptimeRobot to keep the server awake."""
    return {"status": "ok"}


@app.get("/")
async def root():
    """Root endpoint."""
    return {
        "message": "Telegram RAG Summarizer API is running",      
        "endpoints": {
            "health": "/health",
            "admin_feed": "/admin/feed",          
            "webhook": "/webhook"
        }
    }


async def send_message(chat_id: int, text: str):
    """Send a message to the specified chat ID via Telegram API."""
    async with AsyncClient() as client:
        url = f"{TELEGRAM_API_URL}/sendMessage"
        payload = {
            "chat_id": chat_id,
            "text": text,
        }
        response = await client.post(url, json=payload)
        response.raise_for_status()


@app.post("/webhook")
async def webhook(
    request: Request,
    x_telegram_bot_api_secret_token: Optional[str] = Header(None)
):
    """
    Webhook endpoint to receive messages from Telegram.
    Invokes the ChatAgent to perform RAG search and summarization.
    """
   
    print(x_telegram_bot_api_secret_token)
    # 1. Verify webhook secret token (Security)
    if x_telegram_bot_api_secret_token != SECRET_TOKEN:
        raise HTTPException(status_code=403, detail="Unauthorized webhook")

    # 2. Parse the incoming update
    update = await request.json()
    message = update.get("message", {})
    chat_id = message.get("chat", {}).get("id")
    from_user_id = str(message.get("from", {}).get("id"))

    # 3. Authorization Check
    #if from_user_id != MY_CHAT_ID:
    #    return {"status": "ignored", "reason": "unauthorized_user"}

    # 4. Extract text and invoke the Agent
    text = message.get("text")
    if text:
        try:
            # We use 'await' because ChatAgent.run is an async function
            # This triggers the search_docs tool and the Qwen summarization
            print("before invking agent")
            ai_response = await agent.run(text)
            
            # 5. Send the final generated response back to Telegram
            await send_message(chat_id, ai_response)
            
        except Exception as e:
            # Log the error and notify the user so they aren't left waiting
            print(f"Error in Agent Execution: {e}")
            await send_message(chat_id, "I'm sorry, I had trouble accessing my knowledge base. Please try again in a moment.")

    return {"status": "ok"}

@app.post("/admin/feed")
async def admin_feed_knowledge(data: dict = Body(...)):
    """
    HTTP Entry point: Receives request and appends text to knowledge_base.docx.
    Content is appended as a new paragraph with a timestamp separator.
    """
    text_content = data.get("text")

    if not text_content:
        raise HTTPException(status_code=400, detail="No text provided")

    try:
        # Check if the docx file exists
        if os.path.exists(KNOWLEDGE_BASE_PATH):
            # Load existing document (preserving all existing content)
            doc = Document(KNOWLEDGE_BASE_PATH)
            # Add separator with timestamp for new content
            separator = doc.add_paragraph()
            separator.add_run(f"\n--- Added on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} ---\n").bold = True
        else:
            # Create new document
            doc = Document()

        # Add the new text as a paragraph (appends to the end)
        doc.add_paragraph(text_content)

        # Save the document (overwrites the file with all content including new paragraph)
        doc.save(KNOWLEDGE_BASE_PATH)

        return {
            "status": "success",
            "message": "Text appended to knowledge_base.docx"
        }

    except Exception as e:
        print(f"Failed to append to docx: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to append to docx: {str(e)}")




